"""Generate a human-readable Word report from completed run artifacts."""

from __future__ import annotations

import json
import math
import tempfile
from pathlib import Path
from typing import Any, Iterable, Mapping, Optional, Sequence

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "18324A"
TEAL = "247B7B"
GOLD = "C9942E"
PALE_BLUE = "E8F0F5"
PALE_TEAL = "E7F2F1"
PALE_GOLD = "F8F0DF"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "66727D"
WHITE = "FFFFFF"
INK = "1C252C"


def _finite(value: Any) -> bool:
    try:
        return math.isfinite(float(value))
    except (TypeError, ValueError):
        return False


def _number(row: Mapping[str, Any], key: str) -> float:
    value = row.get(key, np.nan)
    return float(value) if _finite(value) else float("nan")


def _fmt(value: Any, digits: int = 2, suffix: str = "") -> str:
    if not _finite(value):
        return "Not reported"
    return f"{float(value):,.{digits}f}{suffix}"


def _first_and_last(summary: pd.DataFrame) -> tuple[pd.Series, pd.Series]:
    if summary.empty:
        raise ValueError("Summary CSV contains no rows")
    ordered = summary.sort_values("time_s", kind="stable") if "time_s" in summary else summary
    return ordered.iloc[0], ordered.iloc[-1]


def _component_names(columns: Iterable[str], prefix: str) -> list[str]:
    return [str(c)[len(prefix) :] for c in columns if str(c).startswith(prefix)]


def _pretty_component(name: str) -> str:
    return name.replace("_", "-").replace("n-", "n-")


def _set_cell_fill(cell: Any, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _set_cell_margins(cell: Any, top: int = 80, start: int = 110, bottom: int = 80, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_table_widths(table: Any, widths: Sequence[float]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)


def _set_run(run: Any, *, size: float = 9.5, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.34)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 26, NAVY, 0, 5),
        ("Heading 1", 16, NAVY, 14, 7),
        ("Heading 2", 12.5, TEAL, 10, 5),
        ("Heading 3", 10.5, NAVY, 7, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Title"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run(run, size=8, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def _add_header_footer(doc: Document, case_name: str) -> None:
    for index, section in enumerate(doc.sections):
        if index > 0:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
        header = section.header.paragraphs[0]
        header.text = "DYNAMIC DISTILLATION  /  RUN REPORT"
        _set_run(header.runs[0], size=8, color=TEAL, bold=True)
        footer = section.footer.paragraphs[0]
        footer.clear()
        left = footer.add_run(case_name[:70])
        _set_run(left, size=8, color=MID_GRAY)
        footer.add_run("\t")
        _add_page_number(footer)


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    widths: Sequence[float],
    *,
    header_fill: str = NAVY,
    font_size: float = 8.5,
) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_widths(table, widths)
    header = table.rows[0]
    _set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        _set_cell_fill(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(text))
        _set_run(run, size=font_size, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                _set_cell_fill(cell, LIGHT_GRAY)
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(value))
            _set_run(run, size=font_size)
    return table


def _add_metric_strip(doc: Document, metrics: Sequence[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    table.style = "Table Grid"
    widths = [6.94 / len(metrics)] * len(metrics)
    _set_table_widths(table, widths)
    for cell, (label, value, fill) in zip(table.rows[0].cells, metrics):
        _set_cell_fill(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label.upper() + "\n")
        _set_run(r1, size=7.5, color=MID_GRAY, bold=True)
        r2 = p.add_run(value)
        _set_run(r2, size=13, color=NAVY, bold=True)


def _plot_series(
    summary: pd.DataFrame,
    output: Path,
    panels: Sequence[tuple[str, Sequence[tuple[str, str, float]]]],
    *,
    cumulative_offset_s: float,
) -> bool:
    if "time_s" not in summary or summary.empty:
        return False
    t_min = (pd.to_numeric(summary["time_s"], errors="coerce") + cumulative_offset_s) / 60.0
    fig, axes = plt.subplots(len(panels), 1, figsize=(9.0, 2.35 * len(panels)), sharex=True)
    if len(panels) == 1:
        axes = [axes]
    palette = ["#247B7B", "#C9942E", "#345B7E", "#A64B3C", "#6C7A3D", "#7D5A8C"]
    plotted = False
    for ax, (title, series) in zip(axes, panels):
        for index, (column, label, scale) in enumerate(series):
            if column not in summary:
                continue
            values = pd.to_numeric(summary[column], errors="coerce") * scale
            if not values.notna().any():
                continue
            ax.plot(t_min, values, label=label, color=palette[index % len(palette)], linewidth=1.8)
            plotted = True
        ax.set_title(title, loc="left", fontsize=10, fontweight="bold", color="#18324A")
        ax.grid(True, alpha=0.22, linewidth=0.7)
        ax.spines[["top", "right"]].set_visible(False)
        ax.tick_params(labelsize=8)
        if ax.lines:
            ax.legend(loc="best", frameon=False, fontsize=8, ncol=min(3, len(ax.lines)))
    axes[-1].set_xlabel("Cumulative simulation time (min)", fontsize=8.5)
    fig.tight_layout(pad=1.1)
    if plotted:
        fig.savefig(output, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return plotted


def _plot_final_profiles(profile: pd.DataFrame, output: Path) -> bool:
    if profile.empty or "time_s" not in profile or "stage" not in profile:
        return False
    t = pd.to_numeric(profile["time_s"], errors="coerce").max()
    final = profile[np.isclose(pd.to_numeric(profile["time_s"], errors="coerce"), t)]
    if "node_type" in final:
        final = final[final["node_type"].astype(str).eq("stage")]
    final = final.sort_values("stage")
    if final.empty:
        return False
    stage = pd.to_numeric(final["stage"], errors="coerce")
    fig, axes = plt.subplots(2, 2, figsize=(9.4, 6.1))
    ax = axes[0, 0]
    ax.plot(stage, pd.to_numeric(final.get("T_F"), errors="coerce"), color="#A64B3C", marker="o", ms=3)
    ax.set_title("Temperature", loc="left", fontweight="bold", color="#18324A")
    ax.set_ylabel("deg F")
    ax = axes[0, 1]
    ax.plot(stage, pd.to_numeric(final.get("P_psia_hyd"), errors="coerce"), color="#345B7E", marker="o", ms=3)
    ax.set_title("Pressure", loc="left", fontweight="bold", color="#18324A")
    ax.set_ylabel("psia")
    ax = axes[1, 0]
    ax.plot(stage, pd.to_numeric(final.get("L_out_used_lbmolph"), errors="coerce"), label="Liquid", color="#247B7B")
    ax.plot(stage, pd.to_numeric(final.get("V_out_lbmolph"), errors="coerce"), label="Vapor", color="#C9942E")
    ax.set_title("Internal traffic", loc="left", fontweight="bold", color="#18324A")
    ax.set_ylabel("lbmol/h")
    ax.legend(frameon=False, fontsize=8)
    ax = axes[1, 1]
    composition_columns = [c for c in final.columns if str(c).startswith("x_") and not str(c).startswith("x_eq_")]
    colors = ["#247B7B", "#C9942E", "#7D5A8C", "#A64B3C"]
    for index, column in enumerate(composition_columns[:4]):
        ax.plot(stage, pd.to_numeric(final[column], errors="coerce"), label=_pretty_component(column[2:]), color=colors[index])
    ax.set_title("Liquid composition", loc="left", fontweight="bold", color="#18324A")
    ax.set_ylabel("mole fraction")
    if ax.lines:
        ax.legend(frameon=False, fontsize=8)
    for ax in axes.flat:
        ax.set_xlabel("Stage")
        ax.grid(True, alpha=0.22, linewidth=0.7)
        ax.spines[["top", "right"]].set_visible(False)
        ax.tick_params(labelsize=8)
    fig.tight_layout(pad=1.2)
    fig.savefig(output, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return True


def _condition_rows(row: Mapping[str, Any]) -> list[list[str]]:
    top_pv = _number(row, "Top_level_ctrl_pv")
    top_sp = _number(row, "Top_level_ctrl_sp")
    bottom_pv = _number(row, "Bottom_level_ctrl_pv")
    bottom_sp = _number(row, "Bottom_level_ctrl_sp")
    top_is_fraction = _finite(top_pv) and _finite(top_sp) and max(abs(top_pv), abs(top_sp)) <= 1.5
    bottom_is_fraction = _finite(bottom_pv) and _finite(bottom_sp) and max(abs(bottom_pv), abs(bottom_sp)) <= 1.5
    return [
        ["Feed", _fmt(_number(row, "F_lbmolph")), "lbmol/h"],
        ["Distillate", _fmt(_number(row, "D_lbmolph")), "lbmol/h"],
        ["Bottoms", _fmt(_number(row, "B_lbmolph")), "lbmol/h"],
        ["Reflux", _fmt(_number(row, "total_reflux_used_lbmolph")), "lbmol/h"],
        ["Top pressure", _fmt(_number(row, "P_top_psia"), 3), "psia"],
        ["Bottom pressure", _fmt(_number(row, "P_bot_psia"), 3), "psia"],
        ["Distillate temperature", _fmt(_number(row, "T_Distillate_F"), 2), "deg F"],
        ["Bottoms temperature", _fmt(_number(row, "T_sump_F"), 2), "deg F"],
        ["Condenser duty", _fmt(_number(row, "Q_cond_used_BTUph") / 1.0e6, 3), "MMBtu/h"],
        ["Reboiler duty", _fmt(_number(row, "Q_reb_used_BTUph") / 1.0e6, 3), "MMBtu/h"],
        [
            "Distillate drum level" if top_is_fraction else "Distillate drum controller PV",
            _fmt(100.0 * top_pv if top_is_fraction else top_pv, 2),
            "%" if top_is_fraction else "lbmol",
        ],
        [
            "Bottoms sump level" if bottom_is_fraction else "Bottoms sump controller PV",
            _fmt(100.0 * bottom_pv if bottom_is_fraction else bottom_pv, 2),
            "%" if bottom_is_fraction else "lbmol",
        ],
        ["Steady-state score", _fmt(_number(row, "steady_state_score"), 4), "-"],
    ]


def _product_composition_rows(row: Mapping[str, Any]) -> list[list[str]]:
    components = sorted(
        set(_component_names(row.keys(), "Distillate_x_")) | set(_component_names(row.keys(), "Bottoms_x_"))
    )
    return [
        [
            _pretty_component(component),
            _fmt(100.0 * _number(row, f"Distillate_x_{component}"), 3, "%"),
            _fmt(100.0 * _number(row, f"Bottoms_x_{component}"), 3, "%"),
        ]
        for component in components
    ]


def _parameter_rows(metadata: Mapping[str, Any]) -> list[list[str]]:
    params = dict(metadata.get("simulation_parameters") or {})
    preferred = [
        "runtime_mode", "thermo_mode", "dwsim_property_package", "integrator", "n_steps", "dt_sec",
        "log_every_n_steps", "thermo_every_n_steps",
        "include_energy", "enable_equilibrium_relaxation", "equilibrium_relaxation_mode", "equilibrium_tau_sec",
        "flash_feed_at_stage_conditions", "enable_liquid_hydraulic_override", "liquid_hydraulic_model",
        "liquid_hydraulic_override_alpha", "enable_level_control", "top_level_pv_mode", "top_level_kc",
        "top_level_ti_sec", "bottom_level_pv_mode", "bottom_level_kc", "bottom_level_ti_sec",
        "enable_pressure_control", "pressure_control_mv", "top_pressure_sp_psia", "top_pressure_kc",
        "top_pressure_ti_sec", "top_pressure_pv_filter_tau_sec", "top_pressure_mv_slew_limit_per_s",
        "condenser_duty_mode", "condenser_duty_btu_per_h", "condenser_duty_min_btu_per_h",
        "condenser_duty_max_btu_per_h", "reboiler_duty_btu_per_h", "vapor_holdup_relaxation_sec",
        "vapor_flow_relaxation_sec", "vapor_flow_zero_temperature_target", "dynamic_vflow_nominal_hi_ratio",
        "steady_state_window_sec", "steady_state_min_time_sec", "steady_state_rel_state_rate_tol_per_s",
        "steady_state_temp_rate_tol_F_per_s", "steady_state_sp_error_tol", "steady_state_require_sp",
    ]
    if not params:
        params = {key: metadata.get(key) for key in ("runtime_mode", "thermo_mode", "n_steps", "dt_sec", "flash_feed_at_stage_conditions")}
    rows: list[list[str]] = []
    used: set[str] = set()
    for key in preferred:
        if key in params and params[key] is not None:
            rows.append([key.replace("_", " ").title(), str(params[key])])
            used.add(key)
    return rows


def generate_run_report(
    metadata_json_path: str | Path,
    *,
    output_path: str | Path | None = None,
    simulation_parameters: Optional[Mapping[str, Any]] = None,
    launch_command: Optional[str] = None,
) -> str:
    """Create a Word report for one completed dynamic run."""
    metadata_path = Path(metadata_json_path).expanduser().resolve()
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    if simulation_parameters is not None:
        metadata["simulation_parameters"] = dict(simulation_parameters)
    if launch_command:
        metadata["launch_command"] = str(launch_command)
    summary_path = Path(str(metadata["summary_csv"])).expanduser()
    profile_path = Path(str(metadata["profile_csv"])).expanduser()
    if not summary_path.is_absolute():
        summary_path = (metadata_path.parent / summary_path).resolve()
    if not profile_path.is_absolute():
        profile_path = (metadata_path.parent / profile_path).resolve()
    summary = pd.read_csv(summary_path)
    profile = pd.read_csv(profile_path)
    start, end = _first_and_last(summary)

    run_id = str(metadata.get("run_id") or metadata_path.stem.replace("run_metadata_", ""))
    case_name = str(metadata.get("run_name") or Path(str(metadata.get("excel_path", "case"))).stem)
    output = Path(output_path).expanduser().resolve() if output_path else metadata_path.parent / f"run_report_{run_id}.docx"
    continuation = bool((metadata.get("native_checkpoint_init") or {}).get("loaded"))
    source_time = float((metadata.get("native_checkpoint_init") or {}).get("source_final_time_s") or 0.0)
    elapsed = float(metadata.get("elapsed_wall_sec") or _number(end, "wall_elapsed_s"))
    sim_time = float(metadata.get("final_time_s") or _number(end, "time_s"))
    sim_wall = sim_time / elapsed if elapsed > 0.0 else float("nan")

    doc = Document()
    _style_document(doc)
    section = doc.sections[0]
    _add_header_footer(doc, case_name)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    _set_run(kicker.add_run("DYNAMIC DISTILLATION  /  COMPLETED RUN"), size=8.5, color=TEAL, bold=True)
    title = doc.add_paragraph(case_name, style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    _set_run(subtitle.add_run(str(metadata.get("run_description") or "Simulation operating report")), size=11, color=MID_GRAY)

    end_score = _number(end, "steady_state_score")
    gate = "PASS" if _number(end, "steady_state_flag") >= 0.5 else "NOT PASSED"
    final_t = pd.to_numeric(profile.get("time_s"), errors="coerce").max()
    final_profile_for_status = profile[
        np.isclose(pd.to_numeric(profile.get("time_s"), errors="coerce"), final_t)
    ]
    flash_failure_values = (
        final_profile_for_status["thermo_flash_failed"]
        if "thermo_flash_failed" in final_profile_for_status
        else pd.Series(0.0, index=final_profile_for_status.index)
    )
    flash_failures = int((pd.to_numeric(flash_failure_values, errors="coerce").fillna(0) > 0.5).sum())
    pressure_error = abs(_number(end, "P_top_psia") - _number(end, "P_top_psia_spec"))
    pressure_off_target = _finite(pressure_error) and pressure_error > max(2.0, 0.01 * _number(end, "P_top_psia_spec"))
    top_level_error = abs(_number(end, "Top_level_ctrl_pv") - _number(end, "Top_level_ctrl_sp"))
    bottom_level_error = abs(_number(end, "Bottom_level_ctrl_pv") - _number(end, "Bottom_level_ctrl_sp"))
    inventory_off_target = (
        (_finite(top_level_error) and top_level_error > 0.05)
        or (_finite(bottom_level_error) and bottom_level_error > 0.05)
    )
    run_validity = "REVIEW" if flash_failures or gate != "PASS" or pressure_off_target or inventory_off_target else "USABLE"
    _add_metric_strip(
        doc,
        [
            ("Run validity", run_validity, PALE_TEAL if run_validity == "USABLE" else PALE_GOLD),
            ("Dynamic gate", gate, PALE_TEAL if gate == "PASS" else PALE_GOLD),
            ("Final score", _fmt(end_score, 3), PALE_BLUE),
            ("Sim / wall", _fmt(sim_wall, 3), LIGHT_GRAY),
        ],
    )

    assessment_notes = []
    if gate != "PASS":
        assessment_notes.append(f"The dynamic acceptance gate did not pass (final score {_fmt(end_score, 3)}).")
    if flash_failures:
        assessment_notes.append(f"The final tray snapshot reports {flash_failures} failed thermo flashes; interpret operating results as invalid until the thermo path is restored.")
    if pressure_off_target:
        assessment_notes.append(f"Top pressure is {_fmt(pressure_error, 2, ' psi')} from its logged target, so the rate gate does not establish acceptable pressure control.")
    if inventory_off_target:
        assessment_notes.append("At least one geometry-based vessel level remains more than five percentage points from setpoint.")
    if max(abs(_number(end, "Top_level_ctrl_pv")), abs(_number(end, "Top_level_ctrl_sp"))) > 1.5:
        assessment_notes.append("The top level loop reported molar holdup rather than a geometry-based level fraction.")
    if max(abs(_number(end, "Bottom_level_ctrl_pv")), abs(_number(end, "Bottom_level_ctrl_sp"))) > 1.5:
        assessment_notes.append("The bottom level loop reported molar holdup rather than a geometry-based level fraction.")
    if assessment_notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        _set_run(p.add_run("AUTOMATED ASSESSMENT  "), size=8.2, color=GOLD, bold=True)
        _set_run(p.add_run(" ".join(assessment_notes)), size=9.2, color=INK)

    doc.add_heading("Run Identity", level=1)
    init_kind = "Continuation from native checkpoint" if continuation else "Fresh start"
    identity = [
        ["Run ID", run_id],
        ["Started", str(metadata.get("started_at_local", "Not reported"))],
        ["Completed", str(metadata.get("ended_at_local", "Not reported"))],
        ["Run type", init_kind],
        ["Input workbook", str(metadata.get("excel_path", "Not reported"))],
        ["Input checkpoint", str((metadata.get("native_checkpoint_init") or {}).get("path") or "None")],
        ["Summary data", str(summary_path)],
        ["Tray-profile data", str(profile_path)],
    ]
    _add_table(doc, ["Item", "Value"], identity, [1.45, 5.49], header_fill=TEAL, font_size=8.2)

    doc.add_heading("Operating Snapshot", level=1)
    start_conditions = _condition_rows(start)
    end_conditions = _condition_rows(end)
    _add_table(
        doc,
        ["Parameter", "Start", "End", "Units"],
        [
            [
                e[0] if s[1] == "Not reported" else s[0],
                s[1],
                e[1],
                e[2] if s[1] == "Not reported" else s[2],
            ]
            for s, e in zip(start_conditions, end_conditions)
        ],
        [2.65, 1.45, 1.45, 1.25],
        font_size=8.3,
    )

    doc.add_heading("Product Composition", level=2)
    start_comps = _product_composition_rows(start)
    end_comps = _product_composition_rows(end)
    comp_rows = []
    for initial, final in zip(start_comps, end_comps):
        comp_rows.append([initial[0], initial[1], final[1], initial[2], final[2]])
    _add_table(
        doc,
        ["Component", "Distillate start", "Distillate end", "Bottoms start", "Bottoms end"],
        comp_rows,
        [1.55, 1.35, 1.35, 1.35, 1.35],
        header_fill=TEAL,
        font_size=8.2,
    )

    doc.add_page_break()
    doc.add_heading("Dynamic Trends", level=1)
    with tempfile.TemporaryDirectory(prefix="distillation_report_") as temp_dir:
        temp = Path(temp_dir)
        flow_chart = temp / "flows.png"
        control_chart = temp / "controls.png"
        profile_chart = temp / "profiles.png"
        top_level_scale = 100.0 if max(
            abs(_number(end, "Top_level_ctrl_pv")), abs(_number(end, "Top_level_ctrl_sp"))
        ) <= 1.5 else 1.0
        bottom_level_scale = 100.0 if max(
            abs(_number(end, "Bottom_level_ctrl_pv")), abs(_number(end, "Bottom_level_ctrl_sp"))
        ) <= 1.5 else 1.0
        level_title = (
            "Vessel levels (%)"
            if top_level_scale == 100.0 and bottom_level_scale == 100.0
            else "Vessel level-controller PVs (mixed units; see operating snapshot)"
        )
        if _plot_series(
            summary,
            flow_chart,
            [
                ("Feed and product flows", (("F_lbmolph", "Feed", 1.0), ("D_lbmolph", "Distillate", 1.0), ("B_lbmolph", "Bottoms", 1.0))),
                ("Column traffic", (("total_reflux_used_lbmolph", "Reflux", 1.0), ("V_condensed_in_lbmolph", "Condensate", 1.0), ("boilup_realized_lbmolph", "Boilup", 1.0))),
            ],
            cumulative_offset_s=source_time,
        ):
            doc.add_picture(str(flow_chart), width=Inches(6.75))
        if _plot_series(
            summary,
            control_chart,
            [
                ("Top pressure", (("P_top_ctrl_pv_psia", "PV", 1.0), ("P_top_psia_spec", "Target", 1.0))),
                (level_title, (("Top_level_ctrl_pv", "Drum PV", top_level_scale), ("Top_level_ctrl_sp", "Drum SP", top_level_scale), ("Bottom_level_ctrl_pv", "Sump PV", bottom_level_scale), ("Bottom_level_ctrl_sp", "Sump SP", bottom_level_scale))),
                ("Heat duties", (("Q_cond_used_BTUph", "Condenser", 1.0e-6), ("Q_reb_used_BTUph", "Reboiler", 1.0e-6))),
                ("Level-controller outputs", (("D_lbmolph", "Distillate MV", 1.0), ("B_lbmolph", "Bottoms MV", 1.0))),
            ],
            cumulative_offset_s=source_time,
        ):
            doc.add_picture(str(control_chart), width=Inches(6.75))

        doc.add_heading("Final Tray Profiles", level=1)
        if _plot_final_profiles(profile, profile_chart):
            doc.add_picture(str(profile_chart), width=Inches(6.75))

        landscape = doc.add_section()
        landscape.orientation = WD_ORIENT.LANDSCAPE
        landscape.page_width = Inches(11.0)
        landscape.page_height = Inches(8.5)
        landscape.top_margin = Inches(0.55)
        landscape.bottom_margin = Inches(0.55)
        landscape.left_margin = Inches(0.55)
        landscape.right_margin = Inches(0.55)
        _add_header_footer(doc, case_name)
        doc.add_heading("Final Tray Profile Table", level=1)
        final_t = pd.to_numeric(profile["time_s"], errors="coerce").max()
        final_profile = profile[np.isclose(pd.to_numeric(profile["time_s"], errors="coerce"), final_t)]
        if "node_type" in final_profile:
            final_profile = final_profile[final_profile["node_type"].astype(str).eq("stage")]
        final_profile = final_profile.sort_values("stage")
        comp_names = [c[2:] for c in final_profile.columns if str(c).startswith("x_") and not str(c).startswith("x_eq_")][:3]
        tray_rows = []
        for _, row in final_profile.iterrows():
            tray_rows.append(
                [
                    str(int(float(row["stage"]))),
                    _fmt(row.get("T_F"), 2),
                    _fmt(row.get("P_psia_hyd"), 3),
                    _fmt(row.get("L_out_used_lbmolph"), 1),
                    _fmt(row.get("V_out_lbmolph"), 1),
                    *[_fmt(row.get(f"x_{name}"), 4) for name in comp_names],
                    *[_fmt(row.get(f"y_{name}"), 4) for name in comp_names],
                ]
            )
        headers = ["Stage", "T (deg F)", "P (psia)", "L (lbmol/h)", "V (lbmol/h)"]
        headers += [f"x {_pretty_component(name)}" for name in comp_names]
        headers += [f"y {_pretty_component(name)}" for name in comp_names]
        widths = [0.55, 0.78, 0.82, 0.95, 0.95] + [0.76] * (2 * len(comp_names))
        _add_table(doc, headers, tray_rows, widths, header_fill=NAVY, font_size=7.2)

    portrait = doc.add_section()
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width = Inches(8.5)
    portrait.page_height = Inches(11.0)
    portrait.top_margin = Inches(0.72)
    portrait.bottom_margin = Inches(0.68)
    portrait.left_margin = Inches(0.78)
    portrait.right_margin = Inches(0.78)
    _add_header_footer(doc, case_name)
    doc.add_heading("Simulation Configuration", level=1)
    _add_table(doc, ["Parameter", "Value"], _parameter_rows(metadata), [3.0, 3.94], header_fill=TEAL, font_size=7.8)
    command = str(metadata.get("launch_command") or "").strip()
    if command:
        doc.add_heading("Exact Launch Command", level=2)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        _set_run(p.add_run(command), size=7.5, color=INK)

    doc.core_properties.title = f"Dynamic Distillation Run Report - {case_name}"
    doc.core_properties.subject = f"Run {run_id}"
    doc.core_properties.keywords = "dynamic distillation, run report, process simulation"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return str(output)


__all__ = ["generate_run_report"]
